import io

import pytest

from packages.doc_parser import DocParseError, parse


def test_parse_markdown_returns_decoded_text():
    md = "# Title\n\nSome body.".encode("utf-8")
    assert parse(md, "spec.md") == "# Title\n\nSome body."


def test_parse_txt_returns_decoded_text():
    assert parse(b"just text", "notes.txt") == "just text"


def test_parse_handles_latin1_fallback():
    raw = "résumé".encode("latin-1")
    out = parse(raw, "doc.md")
    assert "r" in out
    assert "sum" in out


def test_parse_unsupported_extension_raises():
    with pytest.raises(DocParseError) as exc:
        parse(b"...", "spec.docs")
    assert ".docs" in str(exc.value)


def test_parse_docx_extracts_text_and_headings():
    from docx import Document

    doc = Document()
    doc.add_heading("Todo App", level=1)
    doc.add_paragraph("A simple todo list.")
    doc.add_heading("Features", level=2)
    doc.add_paragraph("Create a todo.", style="List Bullet")
    doc.add_paragraph("Empty paragraph below.")
    doc.add_paragraph("")
    doc.add_paragraph("Final line.")

    buf = io.BytesIO()
    doc.save(buf)
    out = parse(buf.getvalue(), "spec.docx")

    assert "# Todo App" in out
    assert "## Features" in out
    assert "A simple todo list." in out
    assert "- Create a todo." in out
    assert "Final line." in out
    assert out.count("\n\n") >= 4


def test_parse_empty_docx_raises():
    from docx import Document

    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    with pytest.raises(DocParseError) as exc:
        parse(buf.getvalue(), "empty.docx")
    assert "no extractable text" in str(exc.value)


def test_parse_invalid_docx_raises():
    with pytest.raises(DocParseError):
        parse(b"not a real docx", "bad.docx")


def test_parse_invalid_pdf_raises():
    with pytest.raises(DocParseError):
        parse(b"not a real pdf", "bad.pdf")


def test_parse_pdf_extracts_text():
    pdf_bytes = _build_minimal_pdf("Hello world\nSecond line.")
    out = parse(pdf_bytes, "spec.pdf")
    assert "Hello world" in out
    assert "Second line." in out


def _build_minimal_pdf(text: str) -> bytes:
    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        ByteStringObject,
        DecodedStreamObject,
        DictionaryObject,
        FloatObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    font = DictionaryObject()
    font[NameObject("/Type")] = NameObject("/Font")
    font[NameObject("/Subtype")] = NameObject("/Type1")
    font[NameObject("/BaseFont")] = NameObject("/Helvetica")
    font_ref = writer._add_object(font)

    resources = page.get("/Resources")
    if not isinstance(resources, DictionaryObject):
        resources = DictionaryObject()
        page[NameObject("/Resources")] = resources
    fonts = DictionaryObject()
    fonts[NameObject("/F1")] = font_ref
    resources[NameObject("/Font")] = fonts

    lines = text.split("\n")
    stream_parts = ["BT", "/F1 14 Tf", "72 720 Td"]
    for i, line in enumerate(lines):
        if i > 0:
            stream_parts.append("0 -18 Td")
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream_parts.append(f"({escaped}) Tj")
    stream_parts.append("ET")
    content_data = "\n".join(stream_parts).encode("latin-1")

    content_stream = DecodedStreamObject()
    content_stream.set_data(content_data)
    content_ref = writer._add_object(content_stream)
    page[NameObject("/Contents")] = content_ref

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()
