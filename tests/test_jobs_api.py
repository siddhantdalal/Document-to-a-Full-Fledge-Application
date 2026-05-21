import json
import time
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from apps.api.main import app
from apps.api.routers import jobs as jobs_module

FIXTURES = Path(__file__).parent / "fixtures"


@pytest.fixture
def client(monkeypatch: pytest.MonkeyPatch) -> TestClient:
    spec = json.loads((FIXTURES / "todo_app.spec.json").read_text())

    from packages.spec_extractor import ExtractionResult

    def fake_extract(_markdown: str, _llm, retries: int = 2, max_tokens_budget=None) -> ExtractionResult:
        usage = {"input": 1200, "output": 800, "total": 2000}
        if max_tokens_budget is not None and usage["total"] > max_tokens_budget:
            from packages.spec_extractor import SpecExtractionError
            raise SpecExtractionError(
                f"Token budget {max_tokens_budget} exceeded ({usage['total']} used)."
            )
        return ExtractionResult(spec=spec, usage=usage)

    class FakeClient:
        def __init__(self, *_args, **_kwargs) -> None:
            pass

    monkeypatch.setattr(jobs_module, "extract_spec", fake_extract)
    monkeypatch.setattr(jobs_module, "AnthropicClient", FakeClient)
    monkeypatch.setattr(jobs_module, "OpenAIClient", FakeClient)
    monkeypatch.setattr(jobs_module, "GeminiClient", FakeClient)
    jobs_module._jobs.clear()
    return TestClient(app)


def _wait_done(client: TestClient, job_id: str, timeout: float = 5.0) -> dict:
    deadline = time.time() + timeout
    while time.time() < deadline:
        res = client.get(f"/jobs/{job_id}")
        body = res.json()
        if body["status"] in ("succeeded", "failed"):
            return body
        time.sleep(0.05)
    raise AssertionError(f"Job {job_id} did not finish within {timeout}s")


def test_post_jobs_returns_pending_with_initial_stages(client: TestClient) -> None:
    with open(FIXTURES / "todo_app.md", "rb") as fh:
        res = client.post(
            "/jobs",
            files={"doc": ("todo_app.md", fh, "text/markdown")},
            data={"provider": "anthropic", "model": "claude-opus-4-7"},
            headers={"X-Provider-Key": "test-key"},
        )
    assert res.status_code == 200
    body = res.json()
    assert body["id"]
    assert [s["name"] for s in body["stages"]] == [
        "extract_spec",
        "generate",
        "validate",
        "reconcile",
        "package",
    ]
    assert body["validation"] is None
    assert body["reconciliation"] is None


def test_pipeline_runs_all_stages_and_serves_zip(client: TestClient) -> None:
    with open(FIXTURES / "todo_app.md", "rb") as fh:
        res = client.post(
            "/jobs",
            files={"doc": ("todo_app.md", fh, "text/markdown")},
            data={"provider": "anthropic", "model": "claude-opus-4-7"},
            headers={"X-Provider-Key": "test-key"},
        )
    job_id = res.json()["id"]
    final = _wait_done(client, job_id)

    assert final["status"] == "succeeded", final.get("error")
    assert final["artifact_ready"] is True
    assert final["spec"]["app"]["name"] == "Todo App"
    assert all(s["status"] == "succeeded" for s in final["stages"])
    for stage in final["stages"]:
        assert stage["started_at"] is not None
        assert stage["finished_at"] is not None
        assert stage["message"]

    assert final["validation"]["ok"] is True
    assert final["validation"]["summary"]["python_files"] > 0
    assert final["usage"] == {"input": 1200, "output": 800, "total": 2000}
    assert final["max_tokens"] is None
    rec = final["reconciliation"]
    assert rec["ok"] is True
    assert rec["coverage"]["entities"]["covered"] == rec["coverage"]["entities"]["total"]
    assert rec["coverage"]["endpoints"]["covered"] == 6
    assert rec["coverage"]["screens"]["covered"] == 3

    artifact = client.get(f"/jobs/{job_id}/artifact")
    assert artifact.status_code == 200
    assert artifact.headers["content-type"] == "application/zip"
    assert "todo-app.zip" in artifact.headers["content-disposition"]
    assert len(artifact.content) > 1000


def test_unknown_job_returns_404(client: TestClient) -> None:
    assert client.get("/jobs/does-not-exist").status_code == 404
    assert client.get("/jobs/does-not-exist/artifact").status_code == 404


def test_artifact_returns_409_while_pending(client: TestClient) -> None:
    jobs_module._jobs["pending-id"] = jobs_module._new_job("pending-id")
    res = client.get("/jobs/pending-id/artifact")
    assert res.status_code == 409


def test_unsupported_provider_returns_400(client: TestClient) -> None:
    with open(FIXTURES / "todo_app.md", "rb") as fh:
        res = client.post(
            "/jobs",
            files={"doc": ("todo_app.md", fh, "text/markdown")},
            data={"provider": "cohere", "model": "command-r"},
            headers={"X-Provider-Key": "test-key"},
        )
    assert res.status_code == 400


def test_pipeline_works_with_openai_provider(client: TestClient) -> None:
    with open(FIXTURES / "todo_app.md", "rb") as fh:
        res = client.post(
            "/jobs",
            files={"doc": ("todo_app.md", fh, "text/markdown")},
            data={"provider": "openai", "model": "gpt-4o"},
            headers={"X-Provider-Key": "sk-test"},
        )
    assert res.status_code == 200
    final = _wait_done(client, res.json()["id"])
    assert final["status"] == "succeeded"


def test_docx_upload_is_parsed_into_pipeline(client: TestClient) -> None:
    import io as _io

    from docx import Document

    doc = Document()
    doc.add_heading("Todo App", level=1)
    doc.add_paragraph("A simple todo app per the docx fixture.")
    buf = _io.BytesIO()
    doc.save(buf)

    res = client.post(
        "/jobs",
        files={"doc": ("spec.docx", buf.getvalue(),
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"provider": "anthropic", "model": "claude-opus-4-7"},
        headers={"X-Provider-Key": "test-key"},
    )
    assert res.status_code == 200
    final = _wait_done(client, res.json()["id"])
    assert final["status"] == "succeeded"


def test_empty_document_returns_400(client: TestClient) -> None:
    res = client.post(
        "/jobs",
        files={"doc": ("blank.md", b"   \n   ", "text/markdown")},
        data={"provider": "anthropic", "model": "claude-opus-4-7"},
        headers={"X-Provider-Key": "test-key"},
    )
    assert res.status_code == 400


def test_pipeline_works_with_gemini_provider(client: TestClient) -> None:
    with open(FIXTURES / "todo_app.md", "rb") as fh:
        res = client.post(
            "/jobs",
            files={"doc": ("todo_app.md", fh, "text/markdown")},
            data={"provider": "gemini", "model": "gemini-2.5-flash"},
            headers={"X-Provider-Key": "g-test"},
        )
    assert res.status_code == 200
    final = _wait_done(client, res.json()["id"])
    assert final["status"] == "succeeded"


def test_token_budget_aborts_pipeline(client: TestClient) -> None:
    with open(FIXTURES / "todo_app.md", "rb") as fh:
        res = client.post(
            "/jobs",
            files={"doc": ("todo_app.md", fh, "text/markdown")},
            data={
                "provider": "anthropic",
                "model": "claude-opus-4-7",
                "max_tokens": "100",
            },
            headers={"X-Provider-Key": "test-key"},
        )
    assert res.status_code == 200
    final = _wait_done(client, res.json()["id"])
    assert final["status"] == "failed"
    assert final["max_tokens"] == 100
    assert "budget" in (final["error"] or "").lower()
    extract_stage = next(s for s in final["stages"] if s["name"] == "extract_spec")
    assert extract_stage["status"] == "failed"


def test_token_budget_must_be_positive(client: TestClient) -> None:
    with open(FIXTURES / "todo_app.md", "rb") as fh:
        res = client.post(
            "/jobs",
            files={"doc": ("todo_app.md", fh, "text/markdown")},
            data={
                "provider": "anthropic",
                "model": "claude-opus-4-7",
                "max_tokens": "0",
            },
            headers={"X-Provider-Key": "test-key"},
        )
    assert res.status_code == 400
